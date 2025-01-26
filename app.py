import os
from flask import Flask, render_template, request, redirect, url_for
from docx import Document
import pandas as pd

app = Flask(__name__)

# Define the UPLOAD_FOLDER path
UPLOAD_FOLDER = os.path.abspath(os.path.join(os.getcwd(), 'uploads'))
app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER

# Create the uploads folder if it doesn't exist
try:
    os.makedirs(UPLOAD_FOLDER, exist_ok=True)
except PermissionError:
    print("Permission denied while creating the uploads folder.")
    # Optionally, you can return an error message here

@app.route("/")
def login():
    return render_template("login.html")

@app.route("/login", methods=["POST"])
def handle_login():
    username = request.form.get("username")
    return redirect(url_for("details", username=username))

@app.route("/details")
def details():
    username = request.args.get("username")
    return render_template("details.html", username=username)

@app.route("/save_details", methods=["POST"])
def save_details():
    program = request.form.get("program")
    section = request.form.get("section")
    year = request.form.get("year")
    academic_year = request.form.get("academicYear")
    semester = request.form.get("semester")
    batch = request.form.get("batch")
    date = request.form.get("date")

    doc = Document("template.docx")
    placeholders = {
        "{{program}}": program,
        "{{sec}}": section,
        "{{year}}": year,
        "{{acyear}}": academic_year,
        "{{sem}}": semester,
        "{{batch}}": batch,
        "{{date}}": date
    }

    for placeholder, value in placeholders.items():
        for paragraph in doc.paragraphs:
            if placeholder in paragraph.text:
                paragraph.text = paragraph.text.replace(placeholder, value)
    
    return redirect(url_for("upload"))

@app.route("/upload")
def upload():
    return render_template("upload.html")

@app.route("/upload_file", methods=["POST"])
def upload_file():
    if 'file' not in request.files:
        return "No file part in the form."
    
    file = request.files['file']
    
    if file.filename == '':
        return "No selected file."

    # Check if the file is an Excel file (Optional, but recommended)
    if not file.filename.endswith('.xlsx'):
        return "Invalid file type. Please upload an Excel file."

    # Save the file to the UPLOAD_FOLDER
    file_path = os.path.join(app.config['UPLOAD_FOLDER'], file.filename)
    
    try:
        file.save(file_path)
    except PermissionError:
        return "Permission denied when saving the file."
    
    # Process the uploaded file after saving it
    return process_file(file_path)

# Function to process the uploaded file (Excel file)
def table1(sub,n,uploaded_file_path):
    global samp
    global num
    global j
    df=pd.read_excel(uploaded_file_path,sheet_name=sub)
    dic1=[]
    dic1.append(n+1)
    dic1.append(sub)
    dic1+=[sub]
    dic1+=[1]
    dic1+=[2]
    dic1+=[3]
    dic1+=[4]
    dic1+=["dfgbn"]
    dic1+=["fds"]
    dic1+=[len(df["name"])]
    dic1+=[len(df[df["UR"]!="ab"])]
    dic1+=[dic1[9]-dic1[10]]
    dic1+=[len(df[(df["UR"] != 'ab') & (df["UR"] != 'RA')])]
    dic1+=[dic1[10]-dic1[12]]
    dic1+=[f"{(dic1[12]/dic1[10])*100:.1f}"]
    if j==0:
        oac=[]
        sheet_names = pd.ExcelFile(uploaded_file_path).sheet_names  # Get sheet names dynamically
        for sheet_name in sheet_names:
            dfs = pd.read_excel(uploaded_file_path, sheet_name=sheet_name) 
            if 'UR' in dfs.columns and 'name' in dfs.columns:
                filtered_names = dfs[(dfs['UR'] != 'ab') & (dfs['UR'] != 'RA')]['name'].tolist()
            oac.append(filtered_names)
        samp=oac[0]
        for sl in range(1,len(oac)):
            samp1=[item for item in samp if item in oac[sl]]
            samp=samp1
        dfs = pd.read_excel(uploaded_file_path, sheet_name="arrear")
        num=len(dfs[dfs["Arrear history"]!=0])
    j=1
    dic1+=[f"{len(samp):.2f}"]
    dic1+=[int((len(samp)/len(df["name"]))*100)]#f"{(len(samp)/len(df["name"]))*100:.2f}"
    dic1+=[num]
    return dic1


def table2(sub,n,uploaded_file_path):
    df=pd.read_excel(uploaded_file_path,sheet_name=sub)
    dic1=[]
    dic1.append(sub)
    dic1+=[sub]
    dic1+=[1]
    dic1+=[2]
    dic1+=[3]
    dic1+=[4]
    dic1+=["dfgbn"]
    dic1+=["fds"]
    dic1+=[len(df["name"]),len(df["name"]),len(df["name"])]*3
    exm=["IA1","IA2","UR"]
    for i in range(2):
        dic1.append(len(df[(df[exm[i]]!="ab") & (df[exm[i]]!=0)]))
    for i in range(3):
        dic1.append(len(df["name"])-len(df[(df[exm[i]]!="ab") & (df[exm[i]]!=0)]))
    for i in range(2):
        dic1.append(len(df[(df[exm[i]] != 0) & (df[exm[i]]>35)]))
    dic1.append(len(df[(df["UR"] != 'ab') & (df["UR"] != 'RA')]))
    for i in range(2):
        dic1.append(len(df[(df[exm[i]] == 0) & (df[exm[i]]<35)]))
    dic1.append(len(df[(df["UR"] == 'ab') & (df["UR"] == 'RA')]))
    for i in range(2):
        dic1.append((len(df[(df[exm[i]] != 0) & (df[exm[i]]>35)])/len(df["name"]))*100)
    dic1.append((len(df[(df["UR"] != 'ab') & (df["UR"] != 'RA')])/len(df["name"]))*100)
    return dic1

def table3(table, uploaded_file_path):
    # Load the "arrear" sheet from the uploaded file
    df = pd.read_excel(uploaded_file_path, sheet_name="arrear")

    # Filter students with arrears
    arrear_students = df[df["Arrear history"] == "ar"]

    # Identify available semester columns dynamically
    semester_columns = [col for col in df.columns if col.startswith("sem")]
    max_semesters = len(semester_columns)

    # Prepare rows for each student
    rows_data = []
    idx = 1
    for idxx, row in arrear_students.iterrows():
        # Calculate total arrears
        total_arrears = sum([row.get(sem, 0) for sem in semester_columns])

        # Create student data dictionary
        student_data = {
            "SL.NO.": idx,
            "REGISTER NUMBER": row.get("regno", "-"),
            "NAME OF THE STUDENT": row.get("name", "-"),
            "NO. OF ARREARS": total_arrears,
            "SEMESTERS": [
                row.get(f"sem{i}", "-") for i in range(1, max_semesters + 1)
            ],  # Fill missing semesters with "-" and ensure it doesn't exceed the available semesters
        }
        idx += 1
        rows_data.append(student_data)

    # Ensure the table has enough rows for the data
    while len(table.rows) < len(rows_data) + 2:  # Add rows after header
        table.add_row()

    # Ensure the table has enough columns to accommodate semester data
    if len(table.rows[0].cells) < 4 + max_semesters:  # 4 fixed columns + number of semester columns
        for row in table.rows:
            while len(row.cells) < 4 + max_semesters:  # Add columns for semesters dynamically
                row.add_cell()

    # Fill the table with data
    for i, student_data in enumerate(rows_data):
        row = table.rows[i + 2]  # Offset for header rows (index 2 is the first data row)
        row.cells[0].text = str(student_data["SL.NO."])
        row.cells[1].text = str(student_data["REGISTER NUMBER"])
        row.cells[2].text = str(student_data["NAME OF THE STUDENT"])
        row.cells[3].text = str(student_data["NO. OF ARREARS"])
        for sem_idx, sem_value in enumerate(student_data["SEMESTERS"]):
            if sem_idx + 4 < len(row.cells):  # Ensure there are enough columns to avoid index issues
                row.cells[4 + sem_idx].text = str(sem_value)

    print("Table 3 has been updated successfully.")


def table4(table, uploaded_file_path):
    # Load Excel file and get sheet names
    xl = pd.ExcelFile(uploaded_file_path)
    sheet_names = xl.sheet_names
    sheet_names.pop()  # Remove last sheet if necessary (adjust as required)

    rows_data = []
    
    # Iterate over each sheet
    for sheet_name in sheet_names:
        # Load the sheet data
        df = xl.parse(sheet_name)

        # Extract semester number (last character of sheet name)
        sem_no = sheet_name[-1] if sheet_name[-1].isdigit() else "N/A"

        # Count arrears and absent
        arrear_students = df[(df["UR"] == "RA") | (df["UR"] == "ab")]
        arrear_count = len(arrear_students)

        # Prepare names and reg. numbers as a single string
        student_list = "\n".join(
            [
                f"{row['name']} ({row['regno']})"
                for _, row in arrear_students.iterrows()
            ]
        )

        # Extract course information
        course_code = sheet_name  # Sheet name as course code
        course_title = df.columns[0] if len(df.columns) > 0 else "N/A"
        L, T, P, C = (1, 2, 3, 4)  # Default values; adjust if needed

        # Prepare data for a row
        rows_data.append({
            "SL.NO.": len(rows_data) + 1,
            "COURSE CODE": course_code,
            "COURSE TITLE": course_title,
            "L": L,
            "T": T,
            "P": P,
            "C": C,
            "SEM NO.": sem_no,
            "NO. OF ARREARS": arrear_count,
            "NAME AND REGISTERNO. OF THE STUDENTS": student_list,
            "ACTION PLAN": "",
            "EXECUTION CONFIRMATION": ""
        })

    # Ensure the table has enough rows for the data
    while len(table.rows) < len(rows_data) + 1:  # Add rows after header
        table.add_row()

    # Fill the table with data
    for i, row_data in enumerate(rows_data):
        row = table.rows[i + 1]  # Offset for header row
        row.cells[0].text = str(row_data["SL.NO."])
        row.cells[1].text = str(row_data["COURSE CODE"])
        row.cells[2].text = str(row_data["COURSE TITLE"])
        row.cells[3].text = str(row_data["L"])
        row.cells[4].text = str(row_data["T"])
        row.cells[5].text = str(row_data["P"])
        row.cells[6].text = str(row_data["C"])
        row.cells[7].text = str(row_data["SEM NO."])
        row.cells[8].text = str(row_data["NO. OF ARREARS"])
        row.cells[9].text = str(row_data["NAME AND REGISTERNO. OF THE STUDENTS"])
        row.cells[10].text = str(row_data["ACTION PLAN"])
        row.cells[11].text = str(row_data["EXECUTION CONFIRMATION"])

    print("Table 4 has been updated successfully.")

def table5(sub, n, uploaded_file_path):
    # Load the data from the uploaded Excel file
    df = pd.read_excel(uploaded_file_path, sheet_name=sub)
    
    # Prepare data for the table
    data = []
    data.append(n)
    data.append(sub)
    data.append(sub)
    data.append("ctype")
    data.append("ctype")
    data.append(1)
    data.append(2)
    data.append(3)
    data.append(4)
    data.append(len(df[df["UR"] == 'O']))
    data.append(len(df[df["UR"] == 'A+']))
    data.append(len(df[df["UR"] == 'A']))
    data.append(len(df[df["UR"] == 'B+']))
    data.append(len(df[df["UR"] == 'B']))
    data.append(len(df[df["UR"] == 'C']))

    print("Table 5 has been updated successfully.")
    return data
def table6(table, uploaded_file_path):
    # Load the arrear data from the uploaded Excel file
    df = pd.read_excel(uploaded_file_path, sheet_name="arrear")
    
    # Sort by CGPA in ascending order and get top 5 rows
    sorted_df = df.sort_values(by="cgpa", ascending=True)
    data = []
    
    # Prepare the data for top 5 students
    for i, (index, row) in enumerate(sorted_df.head(5).iterrows()):
        data.append({
            "SL.NO.": i + 1,
            "REG": row["regno"],
            "NAME": row["name"],
            "CGPA": row["cgpa"],
            "MENTOR": row["mentor"],
        })
    
    # Ensure the table has enough rows for the data
    while len(table.rows) - 2 < len(data):  # Add rows after header
        table.add_row()
    
    # Fill the table with the top 5 students' data
    for i, row_data in enumerate(data):
        row = table.rows[i + 2]  # Offset for header row
        row.cells[0].text = str(row_data["SL.NO."])
        row.cells[1].text = str(row_data["REG"])
        row.cells[2].text = str(row_data["NAME"])
        row.cells[3].text = str(row_data["CGPA"])
        row.cells[4].text = str(row_data["MENTOR"])
    
    print("Table 6 has been updated successfully.")


def table7(table, uploaded_file_path):
    # Load the data from the uploaded Excel file
    df = pd.read_excel(uploaded_file_path, sheet_name="arrear")
    
    data = []
    # Filter rows where "Arrear history" is "ap"
    rwb = df[df["Arrear history"] == "ap"]
    
    i = 0
    # Collect data for each row
    for index, row in rwb.iterrows():
        data.append({
            's': i,
            "rg": row["regno"],
            "sname": row["name"],
            "cg": row.get("cgpa", "-")  # Use "-" if "cgpa" is not available
        })
        i += 1
    
    # Add rows to the table
    while len(table.rows) - 1 < len(data):  # Add rows after the header
        table.add_row()
    
    # Populate the table with data
    for i, student_data in enumerate(data):
        row = table.rows[i + 1]  # Offset for header rows
        row.cells[0].text = str(student_data["s"])
        row.cells[1].text = str(student_data["rg"])
        row.cells[2].text = str(student_data["sname"])
        row.cells[3].text = str(student_data["cg"])
    
    print("Table 7 has been updated successfully.")

def table8(table, uploaded_file_path):
    # Load the data from the uploaded Excel file
    df = pd.read_excel(uploaded_file_path, sheet_name="arrear")
    
    data = []
    # Filter rows where "Arrear history" is "ar" or "ab"
    rwb = df[(df["Arrear history"] == "ar") | (df["Arrear history"] == "ab")]
    
    i = 0
    # Collect data for each row
    for index, row in rwb.iterrows():
        data.append({
            's': i,
            "rg": row["regno"],
            "sname": row["name"],
            "cg": row.get("cgpa", "-")  # Use "-" if "cgpa" is not available
        })
        i += 1
    
    # Add rows to the table
    while len(table.rows) - 1 < len(data):  # Add rows after the header
        table.add_row()
    
    # Populate the table with data
    for i, student_data in enumerate(data):
        row = table.rows[i + 1]  # Offset for header rows
        row.cells[0].text = str(student_data["s"])
        row.cells[1].text = str(student_data["rg"])
        row.cells[2].text = str(student_data["sname"])
        row.cells[3].text = str(student_data["cg"])
    
    print("Table 8 has been updated successfully.")

def artable(excel_path):
    # Load the arrear sheet
    arrear_df = pd.read_excel(excel_path, sheet_name="arrear")

    # Identify semester columns (e.g., columns that contain "sem")
    sem_columns = [col for col in arrear_df.columns if "sem" in col.lower()]
    
    # Add a column for total arrears (sum of all semester arrears)
    arrear_df['total_arrears'] = arrear_df[sem_columns].sum(axis=1)

    # Prepare data structure for 6 arrear categories
    data = [[] for _ in range(6)]

    # Load Excel file and get sheet names (except "arrear" sheet)
    xls = pd.ExcelFile(excel_path)
    sheet_names = xls.sheet_names

    for _, student in arrear_df.iterrows():
        reg_no = str(student["regno"])
        total_arrears = student["total_arrears"]
        
        # Calculate which arrear category the student belongs to
        index = min(total_arrears, 6) - 1
        
        # Only process students with arrears
        if total_arrears > 0:
            # Loop through all course sheets
            for sheet in sheet_names:
                if sheet == "arrear":
                    continue
                
                # Read course data for each sheet
                course_df = pd.read_excel(excel_path, sheet_name=sheet)
                
                # Filter the course dataframe to find the student by regno
                student_row = course_df[course_df["regno"].astype(str) == reg_no]
                
                if not student_row.empty:
                    # Get UR value and check if it's either "AB" or "RA"
                    ur_value = student_row.iloc[0]["UR"].strip().upper()
                    if ur_value in ["AB", "RA"]:
                        # Append student details to the corresponding arrear category (index)
                        data[index].append({
                            "REGISTER NO.": reg_no,
                            "NAME OF THE STUDENT": student["name"],
                            "COURSE CODE": sheet,
                            "COURSE TITLE WITH SEMESTER": f"{sheet[-1]}",  # Assuming course title with semester is just the last digit of the sheet name
                            "COURSE TYPE & LTPC": 1234  # Placeholder, adjust if necessary
                        })
    
    return data


def table17(table,uploaded_file_path):
    xl = pd.ExcelFile(uploaded_file_path)
    sheet_names = xl.sheet_names
    sheet_names.pop()
    rows_data = []
    
    # Iterate over each sheet
    for sheet_name in sheet_names:
        # Load the sheet data
        df = xl.parse(sheet_name)

        # Extract semester number (last character of sheet name)
        sem_no = sheet_name[-1] if sheet_name[-1].isdigit() else "N/A"

        # Count arrears and absent
        arrear_students = df[(df["UR"] == "ab")]
        arrear_count = len(arrear_students)

        # Prepare names and reg. numbers as a single string
        student_list = "\n".join(
            [
                f"{row['name']} ({row['regno']})"
                for _, row in arrear_students.iterrows()
            ]
        )
        reason = "\n".join(
    [
        row['reason for absents'] if pd.notna(row['reason for absents']) and row['reason for absents'].strip() else "not mentioned"
        for _, row in arrear_students.iterrows()
    ])


        # Extract course information
        course_code = sheet_name  # Sheet name as course code
        course_title = df.columns[0] if len(df.columns) > 0 else "N/A"
        L, T, P, C = (1, 2, 3, 4)  # Default values; adjust if needed

        # Prepare data for a row
        rows_data.append({
            "SL.NO.": len(rows_data) + 1,
            "COURSE CODE": course_code,
            "COURSE TITLE": course_title,
            "t":"ty",
            "L":L,
            "T":T,
            "P":P,
            "C":C,
            "NO. OF ABS": arrear_count,
            "NAME AND REGISTERNO. OF THE STUDENTS": student_list,
            "REASON": reason
        })

    # Ensure table has enough rows
    while len(table.rows) < len(rows_data) + 1:  # Add rows after header
        table.add_row()

    # Fill the table
    for i, row_data in enumerate(rows_data):
        row = table.rows[i + 1]  # Offset for header row
        row.cells[0].text = str(row_data["SL.NO."])
        row.cells[1].text = str(row_data["COURSE CODE"])
        row.cells[2].text = str(row_data["COURSE TITLE"])
        row.cells[3].text = str(row_data["L"])
        row.cells[4].text = str(row_data["T"])
        row.cells[5].text = str(row_data["P"])
        row.cells[6].text = str(row_data["C"])
        row.cells[7].text = str(row_data["t"])
        row.cells[8].text = str(row_data["NO. OF ABS"])
        row.cells[9].text = str(row_data["NAME AND REGISTERNO. OF THE STUDENTS"])
        row.cells[10].text = str(row_data["REASON"])

    print("Table 17 has been updated successfully.")

def process_file(file_path):
    # Check if the file exists
    if not os.path.exists(file_path):
        return f"File '{file_path}' not found."
    
    datafile = pd.ExcelFile(file_path)
    
    # Example of working with the Excel file after upload
    subject = ["MA3354", "CS3301", "CS3351", "CS3352", "CS3391"]
    doc = Document("template.docx")
    j = 0
    num = 0
    semester = 4
    samp = []
    
    for table_index, table in enumerate(doc.tables):  # Iterate through each table
        if len(table.rows) > 2 and table_index in [0, 2, 8]:
            if len(table.rows) - 2 < len(subject) and table_index != 4:
                print(len(table.rows), len(subject))
                for _ in range(len(subject) - (len(table.rows) - 2)):
                    table.add_row()  # Add rows

            i = 0
            for row_index, row in enumerate(table.rows):  # Iterate through rows
                if row_index >= 2:  # To ignore the sign table
                    if table_index == 0:
                        for col_in in range(15, 18):
                            start_cell = table.cell(2, col_in)
                            for row_in in range(2, len(table.rows)):
                                start_cell.merge(table.cell(row_in, col_in))
                        if i < len(subject):
                            data = table1(subject[i], i, UPLOAD_FOLDER)
                        flg = 1
                    elif table_index == 2:
                        if i < len(subject):
                            data = table2(subject[i], i, UPLOAD_FOLDER)
                        flg = 1
                    elif table_index == 8:
                        if i < len(subject):
                            data = table5(subject[i], i)
                        flg = 1
                    else:
                        print("else part")
                    
                    if flg == 1:
                        for cell_index, cell in enumerate(row.cells):
                            cell.text = str(data[cell_index])
                        flg = 0
                    i += 1
    
    if table_index == 4:
        table3(table, UPLOAD_FOLDER)
    elif table_index == 6:
        table4(table, UPLOAD_FOLDER)
    elif table_index == 10:
        table6(table)
    elif table_index == 12:
        table7(table)
    elif table_index == 14:
        table8(table)
    elif table_index in [15, 16, 17, 18, 19, 20]:
        if ar == 0:
            data = artable(UPLOAD_FOLDER)
            print(data)
            ar = 1
        while len(table.rows) - 2 < len(data[table_index - 15]):  # Add rows after the header
            table.add_row()
        
        # Populate the table with data
        for i, student_data in enumerate(data[table_index - 15]):
            if len(data[table_index - 15]) == 0:
                break
            row = table.rows[i + 2]  # Offset for header rows
            row.cells[0].text = str(i)
            row.cells[1].text = str(student_data["REGISTER NO."])
            row.cells[2].text = str(student_data["NAME OF THE STUDENT"])
            row.cells[3].text = str(student_data["COURSE CODE"])
            row.cells[4].text = str(student_data["COURSE TITLE WITH SEMESTER"])
            row.cells[5].text = str(student_data["COURSE TYPE & LTPC"])
        print("Table has been updated successfully.")
    
    elif table_index == 23:
        for i in range(len(data)):
            row = table.rows[i + 1]
            row.cells[2].text = str(len(data[i]))
    elif table_index == 24:
        df = pd.read_excel(UPLOAD_FOLDER, sheet_name="arrear")
        data = [0 for i in range(6)]
        for c in df["cgpa"]:
            if c >= 9.50:
                data[0] += 1
            elif c >= 9.00:
                data[1] += 1
            elif c >= 8.50:
                data[2] += 1
            elif c >= 7.50:
                data[3] += 1
            elif c >= 6.50:
                data[4] += 1
            else:
                data[5] += 1
        for i in range(len(data)):
            row = table.rows[i + 1]
            row.cells[2].text = str(data[i])
    elif table_index == 26:
        table17(table, UPLOAD_FOLDER)
    elif table_index == 28:
        pass
    else:
        print(table_index, len(table.rows))

    # Save or do something with the document (e.g., save the Word document)
    doc.save("output.docx")
    
    return "File processed and saved successfully."

if __name__ == "__main__":
    app.run(debug=True)

